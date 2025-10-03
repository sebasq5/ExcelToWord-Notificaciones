# -*- coding: utf-8 -*-
"""
Genera un DOCUMENTO DE WORD (.docx) INDIVIDUAL POR CADA ESTUDIANTE a partir
de una plantilla, llenando los datos desde el archivo Excel 'otro.xlsx'.

Este script crea un archivo .docx separado para cada estudiante, nombrado
con su apellido y nombre para máxima organización.
"""

import pandas as pd
from docx import Document
from pathlib import Path
import os
import sys
import re
import argparse # Usar argparse para argumentos de línea de comandos

# --- CONFIGURACIÓN DE ARCHIVOS ---
BASE_DIR = Path(__file__).parent
EXCEL_FILE = BASE_DIR / "otro.xlsx"
TEMPLATE_DOCX = BASE_DIR / "plantilla.docx"

# --- NOMBRES DE LAS COLUMNAS EN EXCEL ---
COL_CEDULA = "CÉDULA DEL ESTUDIANTE"
COL_ID = "ID ESTUDIANTE"
COL_APELLIDOS = "APELLIDOS"
COL_NOMBRES = "NOMBRES"
COL_CARRERA = "CARRERA"
COL_TEMA = "TEMA"
COL_TRIB1 = "TRL1"
COL_TRIB2 = "TRL2"
COL_TRIB3 = "TRL3"

def reemplazar_texto(doc, replacements):
    """Busca y reemplaza los marcadores de texto en el documento."""
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(key, str(value))

def main(output_dir):
    """Función principal del script."""
    # 1. VERIFICAR ARCHIVOS DE ENTRADA
    if not os.path.exists(EXCEL_FILE):
        print(f"Error Crítico: No se encuentra el archivo Excel: {EXCEL_FILE}")
        sys.exit(1)
    if not os.path.exists(TEMPLATE_DOCX):
        print(f"Error Crítico: No se encuentra la plantilla de Word: {TEMPLATE_DOCX}")
        sys.exit(1)

    # Crear la carpeta de salida si no existe
    output_dir.mkdir(parents=True, exist_ok=True)

    # 2. LEER EL ARCHIVO EXCEL
    try:
        df = pd.read_excel(EXCEL_FILE, header=0)
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"Error al leer el archivo Excel: {e}")
        sys.exit(1)

    # 3. VERIFICAR COLUMNAS
    required_cols = [COL_CEDULA, COL_ID, COL_APELLIDOS, COL_NOMBRES, COL_CARRERA, COL_TEMA, COL_TRIB1, COL_TRIB2, COL_TRIB3]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        print(f"Error: Faltan las siguientes columnas en 'otro.xlsx': {missing_cols}")
        sys.exit(1)

    # 4. PROCESAR CADA ESTUDIANTE
    filas_validas = df.dropna(subset=[COL_CEDULA]).copy()
    if filas_validas.empty:
        print("Advertencia: No se encontraron estudiantes con cédula.")
        return

    print(f"Se encontraron {len(filas_validas)} estudiantes. Generando documentos individuales...")

    for index, row in filas_validas.iterrows():
        replacements = {
            "{{NOMBRE_COMPLETO}}": f"{row.get(COL_NOMBRES, '')} {row.get(COL_APELLIDOS, '')}".strip(),
            "{{CEDULA}}": str(row.get(COL_CEDULA, '')).strip(),
            "{{TEMA}}": str(row.get(COL_TEMA, '')).strip().capitalize(),
            "{{ID}}": str(row.get(COL_ID, '')).strip(),
            "{{CARRERA}}": str(row.get(COL_CARRERA, '')).strip(),
            "{{TRIBUNAL_1}}": str(row.get(COL_TRIB1, '')) if pd.notna(row.get(COL_TRIB1)) else "",
            "{{TRIBUNAL_2}}": str(row.get(COL_TRIB2, '')) if pd.notna(row.get(COL_TRIB2)) else "",
            "{{TRIBUNAL_3}}": str(row.get(COL_TRIB3, '')) if pd.notna(row.get(COL_TRIB3)) else "",
        }

        doc = Document(TEMPLATE_DOCX)
        reemplazar_texto(doc, replacements)

        nombre_archivo = f"Notificacion_{replacements['{{NOMBRE_COMPLETO}}']}.docx"
        nombre_archivo_seguro = re.sub(r'[^a-zA-Z0-9_\.]', '_', nombre_archivo)
        ruta_salida = output_dir / nombre_archivo_seguro

        try:
            doc.save(ruta_salida)
            print(f"- Generado: {nombre_archivo_seguro}")
        except Exception as e:
            print(f"- ERROR al guardar para {replacements['{{NOMBRE_COMPLETO}}']}: {e}")

    print(f"\n¡ÉXITO! Proceso completado.")
    print(f"Se han guardado {len(filas_validas)} documentos en la carpeta: {output_dir.resolve()}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Genera documentos de Word individuales para notificaciones a estudiantes.")
    parser.add_argument(
        "-o", "--output_dir",
        default=BASE_DIR / "Notificaciones_Generadas",
        help="Directorio donde se guardarán los archivos generados. Por defecto, 'Notificaciones_Generadas'."
    )
    args = parser.parse_args()
    main(Path(args.output_dir))