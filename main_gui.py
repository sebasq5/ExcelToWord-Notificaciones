# -*- coding: utf-8 -*-
"""
Interfaz gráfica de usuario (GUI) para el generador de notificaciones de Word.

Permite al usuario elegir entre generar archivos individuales por estudiante
o un único documento consolidado con todos los estudiantes.
"""

import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog
import pandas as pd
from docx import Document
from pathlib import Path
import os
import sys
import re
import threading

# --- CONFIGURACIÓN GLOBAL ---

def resource_path(relative_path: str) -> Path:
    """
    Obtiene la ruta absoluta a un recurso. Funciona para el modo de desarrollo
    y para el ejecutable de PyInstaller.
    """
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller crea una carpeta temporal y guarda la ruta en _MEIPASS
        base_path = Path(sys._MEIPASS)
    else:
        # No estamos en un paquete, así que la base es el directorio del script
        base_path = Path(__file__).parent

    return base_path / relative_path

EXCEL_FILE = resource_path("otro.xlsx")
TEMPLATE_DOCX = resource_path("plantilla.docx")


# --- NOMBRES DE COLUMNAS (igual que en los otros scripts) ---
COL_CEDULA = "CÉDULA DEL ESTUDIANTE"
COL_ID = "ID ESTUDIANTE"
COL_APELLIDOS = "APELLIDOS"
COL_NOMBRES = "NOMBRES"
COL_CARRERA = "CARRERA"
COL_TEMA = "TEMA"
COL_TRIB1 = "TRL1"
COL_TRIB2 = "TRL2"
COL_TRIB3 = "TRL3"

# --- LÓGICA DE GENERACIÓN (funciones adaptadas de los scripts anteriores) ---

def reemplazar_texto(doc, replacements):
    """Busca y reemplaza texto en párrafos y tablas."""
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(key, str(value))

def leer_y_validar_excel(log_area):
    """Lee y valida el archivo Excel. Devuelve el DataFrame o None si hay error."""
    if not os.path.exists(EXCEL_FILE):
        log_area.insert(tk.END, f"Error Crítico: No se encuentra el archivo Excel: {EXCEL_FILE}\n")
        return None
    if not os.path.exists(TEMPLATE_DOCX):
        log_area.insert(tk.END, f"Error Crítico: No se encuentra la plantilla de Word: {TEMPLATE_DOCX}\n")
        return None

    try:
        df = pd.read_excel(EXCEL_FILE, header=0)
        df.columns = df.columns.str.strip()
        log_area.insert(tk.END, "Archivo Excel leído correctamente.\n")
    except Exception as e:
        log_area.insert(tk.END, f"Error al leer el archivo Excel: {e}\n")
        return None

    required_cols = [COL_CEDULA, COL_ID, COL_APELLIDOS, COL_NOMBRES, COL_CARRERA, COL_TEMA, COL_TRIB1, COL_TRIB2, COL_TRIB3]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        log_area.insert(tk.END, f"Error: Faltan columnas en 'otro.xlsx': {missing_cols}\n")
        return None

    filas_validas = df.dropna(subset=[COL_CEDULA]).copy()
    if filas_validas.empty:
        log_area.insert(tk.END, "Advertencia: No se encontraron estudiantes con cédula.\n")
        return None
        
    return filas_validas

def generar_individuales(log_area, root, output_path_var):
    """Lógica para generar archivos de Word individuales."""
    output_dir_str = output_path_var.get()
    if not output_dir_str or not os.path.isdir(output_dir_str):
        log_area.delete('1.0', tk.END)
        log_area.insert(tk.END, "Error: Por favor, seleccione una carpeta de destino válida primero.\n")
        return

    output_dir = Path(output_dir_str)
    log_area.delete('1.0', tk.END)
    log_area.insert(tk.END, f"Iniciando generación de archivos individuales en: {output_dir}\n")
    root.update_idletasks()

    filas_validas = leer_y_validar_excel(log_area)
    if filas_validas is None:
        log_area.insert(tk.END, "Proceso detenido por errores.\n")
        return

    log_area.insert(tk.END, f"Se encontraron {len(filas_validas)} estudiantes.\n")

    for _, row in filas_validas.iterrows():
        replacements = {
            "{{NOMBRE_COMPLETO}}": f"{row.get(COL_NOMBRES, '')} {row.get(COL_APELLIDOS, '')}".strip(),
            "{{CEDULA}}": str(row.get(COL_CEDULA, '')).strip(),
            "{{TEMA}}": str(row.get(COL_TEMA, '')).strip().capitalize(),
            "{{ID}}": str(row.get(COL_ID, '')).strip(),
            "{{CARRERA}}": str(row.get(COL_CARRERA, '')).strip(),
            "{{TRIBUNAL_1}}": str(row.get(COL_TRIB1, '')) if pd.notna(row.get(COL_TRIB1)) else "",
            "{{TRIBUNAL_2}}": str(row.get(COL_TRIB2, '')) if pd.notna(row.get(COL_TRIB2)) else "",
            "{{TRIBUNAL_3}}": str(row.get(COL_TRIB3, '')) if pd.notna(row.get(COL_TRIB3)) else "",
        }
        doc = Document(TEMPLATE_DOCX)
        reemplazar_texto(doc, replacements)
        nombre_archivo = f"Notificacion_{replacements['{{NOMBRE_COMPLETO}}']}.docx"
        nombre_archivo_seguro = re.sub(r'[^a-zA-Z0-9_\.]', '_', nombre_archivo)
        ruta_salida = output_dir / nombre_archivo_seguro
        doc.save(ruta_salida)
        log_area.insert(tk.END, f"- Generado: {nombre_archivo_seguro}\n")
        root.update_idletasks()

    log_area.insert(tk.END, f"\n¡ÉXITO! Proceso completado.\nSe guardaron los archivos en: {output_dir.resolve()}\n")

def generar_unico(log_area, root, output_path_var):
    """Lógica para generar un único archivo de Word con saltos de página."""
    output_dir_str = output_path_var.get()
    if not output_dir_str or not os.path.isdir(output_dir_str):
        log_area.delete('1.0', tk.END)
        log_area.insert(tk.END, "Error: Por favor, seleccione una carpeta de destino válida primero.\n")
        return
        
    output_dir = Path(output_dir_str)
    log_area.delete('1.0', tk.END)
    log_area.insert(tk.END, f"Iniciando generación de documento único en: {output_dir}\n")
    root.update_idletasks()

    filas_validas = leer_y_validar_excel(log_area)
    if filas_validas is None:
        log_area.insert(tk.END, "Proceso detenido por errores.\n")
        return

    log_area.insert(tk.END, f"Se encontraron {len(filas_validas)} estudiantes.\n")

    # --- Procesa el primer estudiante para crear el documento base ---
    primer_estudiante = filas_validas.iloc[0]
    replacements = {
        "{{NOMBRE_COMPLETO}}": f"{primer_estudiante.get(COL_NOMBRES, '')} {primer_estudiante.get(COL_APELLIDOS, '')}".strip(),
        "{{CEDULA}}": str(primer_estudiante.get(COL_CEDULA, '')).strip(),
        "{{TEMA}}": str(primer_estudiante.get(COL_TEMA, '')).strip().capitalize(),
        "{{ID}}": str(primer_estudiante.get(COL_ID, '')).strip(),
        "{{CARRERA}}": str(primer_estudiante.get(COL_CARRERA, '')).strip(),
        "{{TRIBUNAL_1}}": str(primer_estudiante.get(COL_TRIB1, '')) if pd.notna(primer_estudiante.get(COL_TRIB1)) else "",
        "{{TRIBUNAL_2}}": str(primer_estudiante.get(COL_TRIB2, '')) if pd.notna(primer_estudiante.get(COL_TRIB2)) else "",
        "{{TRIBUNAL_3}}": str(primer_estudiante.get(COL_TRIB3, '')) if pd.notna(primer_estudiante.get(COL_TRIB3)) else "",
    }
    doc_final = Document(TEMPLATE_DOCX)
    reemplazar_texto(doc_final, replacements)
    log_area.insert(tk.END, f"- Procesado: {replacements['{{NOMBRE_COMPLETO}}']}\n")
    root.update_idletasks()

    # --- Procesa el resto de los estudiantes ---
    if len(filas_validas) > 1:
        for _, row in filas_validas.iloc[1:].iterrows():
            # Añade un salto de página ANTES de agregar el nuevo contenido
            doc_final.add_page_break()
            
            replacements = {
                "{{NOMBRE_COMPLETO}}": f"{row.get(COL_NOMBRES, '')} {row.get(COL_APELLIDOS, '')}".strip(),
                "{{CEDULA}}": str(row.get(COL_CEDULA, '')).strip(),
                "{{TEMA}}": str(row.get(COL_TEMA, '')).strip().capitalize(),
                "{{ID}}": str(row.get(COL_ID, '')).strip(),
                "{{CARRERA}}": str(row.get(COL_CARRERA, '')).strip(),
                "{{TRIBUNAL_1}}": str(row.get(COL_TRIB1, '')) if pd.notna(row.get(COL_TRIB1)) else "",
                "{{TRIBUNAL_2}}": str(row.get(COL_TRIB2, '')) if pd.notna(row.get(COL_TRIB2)) else "",
                "{{TRIBUNAL_3}}": str(row.get(COL_TRIB3, '')) if pd.notna(row.get(COL_TRIB3)) else "",
            }
            
            # Crea un documento temporal para rellenar la plantilla
            doc_template = Document(TEMPLATE_DOCX)
            reemplazar_texto(doc_template, replacements)
            
            # Añade el contenido del temporal al documento final
            for element in doc_template.element.body:
                doc_final.element.body.append(element)
            
            log_area.insert(tk.END, f"- Procesado: {replacements['{{NOMBRE_COMPLETO}}']}\n")
            root.update_idletasks()

    output_path = output_dir / "notificaciones_TODOS_EN_UNO.docx"
    doc_final.save(output_path)
    log_area.insert(tk.END, f"\n¡ÉXITO! Proceso completado.\nDocumento único guardado en: {output_path.resolve()}\n")

# --- CREACIÓN DE LA INTERFAZ GRÁFICA ---
def main_gui():
    root = tk.Tk()
    root.title("Generador de Notificaciones v1.1")
    root.geometry("600x500")
    root.resizable(False, False)

    style = ttk.Style()
    style.configure("TButton", font=("Helvetica", 12), padding=10)
    style.configure("TLabel", font=("Helvetica", 14, "bold"))
    style.configure("TFrame", background="#f0f0f0")

    main_frame = ttk.Frame(root, padding="20 20 20 20")
    main_frame.pack(expand=True, fill=tk.BOTH)

    title_label = ttk.Label(main_frame, text="Generador de Notificaciones")
    title_label.pack(pady=(0, 15))

    # --- Selección de Directorio ---
    output_path_var = tk.StringVar()
    
    def select_output_dir():
        path = filedialog.askdirectory(title="Seleccione dónde guardar los archivos")
        if path:
            output_path_var.set(path)
            dir_label.config(text=path, foreground="black")
            log_area.insert(tk.END, f"Carpeta de destino: {path}\n")
        else:
            log_area.insert(tk.END, "Selección de carpeta cancelada.\n")

    dir_frame = ttk.Frame(main_frame)
    dir_frame.pack(fill=tk.X, pady=10)
    
    btn_select_dir = ttk.Button(
        dir_frame,
        text="Seleccionar Carpeta de Destino",
        command=select_output_dir
    )
    btn_select_dir.pack(side=tk.LEFT, expand=True, fill=tk.X)
    
    dir_label = ttk.Label(main_frame, text="Aún no ha seleccionado una carpeta...", font=("Helvetica", 9), foreground="grey")
    dir_label.pack(fill=tk.X, pady=(0, 10))

    # --- Botones de Acción ---
    btn_individual = ttk.Button(
        main_frame, 
        text="1. Generar Archivos Individuales", 
        command=lambda: threading.Thread(target=generar_individuales, args=(log_area, root, output_path_var)).start()
    )
    btn_individual.pack(fill=tk.X, pady=5)

    btn_unico = ttk.Button(
        main_frame, 
        text="2. Generar Documento Único", 
        command=lambda: threading.Thread(target=generar_unico, args=(log_area, root, output_path_var)).start()
    )
    btn_unico.pack(fill=tk.X, pady=5)

    # --- Área de logs ---
    log_label = ttk.Label(main_frame, text="Progreso:", font=("Helvetica", 10))
    log_label.pack(pady=(15, 5), anchor="w")
    log_area = scrolledtext.ScrolledText(main_frame, height=10, wrap=tk.WORD, font=("Courier New", 9))
    log_area.pack(expand=True, fill=tk.BOTH)

    root.mainloop()

if __name__ == "__main__":
    main_gui()