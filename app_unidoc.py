# -*- coding: utf-8 -*-
"""
Genera un ÚNICO documento Word (.docx) con todos los estudiantes,
cada uno en su propia página, usando 'plantilla.docx' y 'otro.xlsx'.
"""

from pathlib import Path
import sys, os, re, argparse
import pandas as pd
from docx import Document

# --- RUTAS DE ARCHIVOS ---
BASE_DIR = Path(__file__).parent
EXCEL_FILE    = BASE_DIR / "otro.xlsx"
TEMPLATE_DOCX = BASE_DIR / "plantilla.docx"

# --- NOMBRES DE COLUMNAS ---
COL_CEDULA    = "CÉDULA DEL ESTUDIANTE"
COL_ID        = "ID ESTUDIANTE"
COL_APELLIDOS = "APELLIDOS"
COL_NOMBRES   = "NOMBRES"
COL_CARRERA   = "CARRERA"
COL_TEMA      = "TEMA"
COL_TRIB1     = "TRL1"
COL_TRIB2     = "TRL2"
COL_TRIB3     = "TRL3"

# --- MARCADORES EN LA PLANTILLA ---
MARKERS = {
    "{{NOMBRE_COMPLETO}}",
    "{{CEDULA}}",
    "{{TEMA}}",
    "{{ID}}",
    "{{CARRERA}}",
    "{{TRIBUNAL_1}}",
    "{{TRIBUNAL_2}}",
    "{{TRIBUNAL_3}}",
}

# ---------- UTILIDADES ----------
def merge_runs(p):
    """Une todos los runs de un párrafo para que un marcador no quede partido."""
    if len(p.runs) <= 1:
        return
    txt = "".join(run.text for run in p.runs)
    for _ in range(len(p.runs) - 1):
        p.runs[0]._element.getparent().remove(p.runs[1]._element)
    p.runs[0].text = txt


def replace_markers(doc: Document, mapping: dict):
    """Reemplaza marcador → valor en todo el documento (párrafos + tablas)."""
    # Párrafos
    for p in doc.paragraphs:
        merge_runs(p)
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    # Tablas
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    merge_runs(p)
                    for k, v in mapping.items():
                        if k in p.text:
                            p.text = p.text.replace(k, v)
# ----------------------------------


def main(output_path: Path):
    # 1. Validación de archivos
    for f, tag in ((EXCEL_FILE, "Excel"), (TEMPLATE_DOCX, "plantilla")):
        if not f.exists():
            sys.exit(f"❌ No se encuentra el archivo {tag}: {f}")

    # 2. Leer Excel
    try:
        df = pd.read_excel(EXCEL_FILE, header=0).copy()
        df.columns = df.columns.str.strip()
    except Exception as e:
        sys.exit(f"❌ Error al leer Excel: {e}")

    # 3. Verificar columnas obligatorias
    required = [
        COL_CEDULA, COL_ID, COL_APELLIDOS, COL_NOMBRES,
        COL_CARRERA, COL_TEMA, COL_TRIB1, COL_TRIB2, COL_TRIB3
    ]
    missing = [c for c in required if c not in df.columns]
    if missing:
        sys.exit(f"❌ Faltan columnas en Excel: {missing}")

    # 4. Filtrar estudiantes con cédula
    estudiantes = df.dropna(subset=[COL_CEDULA])
    if estudiantes.empty:
        sys.exit("⚠️  No se encontraron estudiantes con cédula.")

    print(f"➡️  Generando documento para {len(estudiantes)} estudiantes…")

    # 5. Documento final
    doc_final = Document()

    for idx, row in estudiantes.iterrows():
        # a) Crear copia de la plantilla
        doc_tmp = Document(TEMPLATE_DOCX)

        # b) Diccionario de reemplazos
        ctx = {
            "{{NOMBRE_COMPLETO}}": f"{row[COL_NOMBRES]} {row[COL_APELLIDOS]}".strip(),
            "{{CEDULA}}":          str(row[COL_CEDULA]).strip(),
            "{{TEMA}}":            str(row[COL_TEMA]).strip().capitalize(),
            "{{ID}}":              str(row[COL_ID]).strip(),
            "{{CARRERA}}":         str(row[COL_CARRERA]).strip(),
            "{{TRIBUNAL_1}}":      str(row[COL_TRIB1]) if pd.notna(row[COL_TRIB1]) else "",
            "{{TRIBUNAL_2}}":      str(row[COL_TRIB2]) if pd.notna(row[COL_TRIB2]) else "",
            "{{TRIBUNAL_3}}":      str(row[COL_TRIB3]) if pd.notna(row[COL_TRIB3]) else "",
        }

        # c) Reemplazar marcadores
        replace_markers(doc_tmp, ctx)

        # d) Añadir al documento final
        if idx > 0:            # ⇦ salto SOLO antes de los siguientes
            doc_final.add_page_break()
        for el in doc_tmp.element.body:
            doc_final.element.body.append(el)

        print(f"   ✔ {ctx['{{NOMBRE_COMPLETO}}']}")

    # 6. Guardar
    try:
        doc_final.save(output_path)
        print(f"\n✅ Documento creado: {output_path.resolve()}")
    except Exception as e:
        sys.exit(f"❌ Error al guardar DOCX: {e}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Genera un único docx con notificaciones para todos los estudiantes."
    )
    parser.add_argument(
        "-o", "--output",
        default=BASE_DIR / "notificaciones_TODOS_EN_UNO.docx",
        help="Ruta del archivo DOCX de salida (por defecto, en el directorio del script)."
    )
    args = parser.parse_args()
    main(Path(args.output))
